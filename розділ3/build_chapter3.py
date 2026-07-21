#!/usr/bin/env python3
"""Build Chapter 3 of the dissertation as a .docx file."""

from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

OUTPUT_PATH = "/sessions/exciting-eloquent-clarke/mnt/КЛОД/розділ3/Розділ_3.docx"


def set_default_font(document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    rpr = style.element.rPr
    if rpr is not None and rpr.rFonts is not None:
        rpr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def set_section_format(section):
    section.page_width = Mm(215.9)
    section.page_height = Mm(279.4)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.0)


def add_chapter_heading(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_section_heading(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_subsection_heading(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_paragraph(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_caption(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_table_simple(document, headers, rows, col_widths_cm=None):
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    for ri, row in enumerate(rows, start=1):
        rcells = table.rows[ri].cells
        for ci, val in enumerate(row):
            rcells[ci].text = ""
            p = rcells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[ci].width = Cm(w)


def page_break(document):
    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


CONTENT = []
def H1(t): CONTENT.append(("h1", t))
def H2(t): CONTENT.append(("h2", t))
def H3(t): CONTENT.append(("h3", t))
def P(t):  CONTENT.append(("p", t))
def CAP(t):CONTENT.append(("cap", t))
def TBL(headers, rows, col_widths_cm=None):
    CONTENT.append(("table", headers, rows, col_widths_cm))
def PB(): CONTENT.append(("pb",))


# Chapter content will be appended via separate file
import sys, os
sys.path.insert(0, os.path.dirname(__file__))
from chapter3_text import build
build(H1, H2, H3, P, CAP, TBL, PB)


def render(document):
    for item in CONTENT:
        if item[0] == "h1":
            add_chapter_heading(document, item[1])
        elif item[0] == "h2":
            add_section_heading(document, item[1])
        elif item[0] == "h3":
            add_subsection_heading(document, item[1])
        elif item[0] == "p":
            add_paragraph(document, item[1])
        elif item[0] == "cap":
            add_caption(document, item[1])
        elif item[0] == "table":
            add_table_simple(document, item[1], item[2], item[3])
        elif item[0] == "pb":
            page_break(document)


def main():
    doc = Document()
    set_default_font(doc)
    for sec in doc.sections:
        set_section_format(sec)
    render(doc)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")
    total_words = 0
    for item in CONTENT:
        if item[0] in ("p", "h1", "h2", "h3", "cap"):
            total_words += len(item[1].split())
    print(f"Total words (approx): {total_words}")


if __name__ == "__main__":
    main()
