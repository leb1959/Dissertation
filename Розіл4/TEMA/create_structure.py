from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

# Кольори
GREEN  = (0x00, 0x70, 0x00)
ORANGE = (0xC0, 0x60, 0x00)
BLUE   = (0x18, 0x5F, 0xA5)
GRAY   = (0x50, 0x50, 0x50)

BG_GREEN  = 'E2EFDA'
BG_ORANGE = 'FFF2CC'
BG_BLUE   = 'DDEEFF'
BG_WHITE  = 'FFFFFF'

def set_cell_bg(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)

def separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '4')
    b.set(qn('w:space'), '1')
    b.set(qn('w:color'), 'CCCCCC')
    pBdr.append(b)
    pPr.append(pBdr)

def chapter_row(table, num, title, status, bg, color, badge_text, badge_bg):
    row = table.add_row()
    # Номер
    c0 = row.cells[0]
    set_cell_bg(c0, bg)
    p = c0.paragraphs[0]
    r = p.add_run(num)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color)

    # Назва
    c1 = row.cells[1]
    set_cell_bg(c1, bg)
    p = c1.paragraphs[0]
    r = p.add_run(title)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color)

    # Статус
    c2 = row.cells[2]
    set_cell_bg(c2, badge_bg)
    p = c2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(badge_text)
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color)

    return row

def sub_row(table, num, title, color, bg, note=None):
    row = table.add_row()
    c0 = row.cells[0]
    set_cell_bg(c0, bg)
    p = c0.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(num)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(*GRAY)

    c1 = row.cells[1]
    set_cell_bg(c1, bg)
    p = c1.paragraphs[0]
    r = p.add_run(title)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(*color)
    if color != GRAY:
        r.font.bold = True

    if note:
        c2 = row.cells[2]
        set_cell_bg(c2, bg)
        p = c2.paragraphs[0]
        r = p.add_run(note)
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    else:
        set_cell_bg(row.cells[2], bg)

    return row

# ============================================================
# TITLE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ПОВНА СТРУКТУРА ДИСЕРТАЦІЇ')
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
p.paragraph_format.space_after = Pt(4)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('"Підвищення екологічної безпеки моніторингу атмосферного повітря шляхом програмно-алгоритмічної корекції індикативних вимірювань PM₁₀ за синхронізованими даними TSP"')
r2.font.size = Pt(11)
r2.font.italic = True
r2.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
p2.paragraph_format.space_after = Pt(12)

# Легенда
legend_table = doc.add_table(rows=1, cols=3)
legend_table.style = 'Table Grid'
legend_data = [
    ('■  Написано — без змін', BG_GREEN, GREEN),
    ('■  Написано — потрібні правки', BG_ORANGE, ORANGE),
    ('■  Новий розділ — створити', BG_BLUE, BLUE),
]
for i, (text, bg, color) in enumerate(legend_data):
    cell = legend_table.rows[0].cells[i]
    set_cell_bg(cell, bg)
    cell.width = Cm(5.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color)

doc.add_paragraph()

# ============================================================
# MAIN TABLE
# ============================================================
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

# Header
hrow = table.rows[0]
for i, (txt, w) in enumerate([('№', 1.5), ('Розділ / підрозділ', 12), ('Статус / примітка', 3)]):
    c = hrow.cells[i]
    c.width = Cm(w)
    set_cell_bg(c, '1F497D')
    p = c.paragraphs[0]
    r = p.add_run(txt)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ---- ВСТУП ----
chapter_row(table, '', 'ВСТУП', 'fix', BG_ORANGE, ORANGE, 'правки', BG_ORANGE)
sub_row(table, '', 'Назва — замінити на нову (16 слів)', ORANGE, BG_ORANGE, 'нова назва затверджена')
sub_row(table, '', 'Мета — уточнити: RF_M1 + RF_M2', ORANGE, BG_ORANGE)
sub_row(table, '', 'Наукова новизна — додати діагностику зональності + програмний комплекс', ORANGE, BG_ORANGE)
sub_row(table, '', 'Практичне значення — ~120 постів Гідромету по Україні', ORANGE, BG_ORANGE)
sub_row(table, '', 'Прибрати посилання на важкі метали', ORANGE, BG_ORANGE)

# ---- Р1 ----
chapter_row(table, 'Р 1', 'Теоретичні основи та нормативно-правова база', 'ok', BG_GREEN, GREEN, 'без змін', BG_GREEN)
sub_row(table, '1.1', 'Екологічна безпека у сфері атмосферного повітря', GRAY, BG_WHITE)
sub_row(table, '1.2', 'Нормативно-правове забезпечення (КМУ №827, Директива 2008/50/EC, №513/2024, №590)', GRAY, BG_WHITE)
sub_row(table, '1.3', 'Референтні методи вимірювання', GRAY, BG_WHITE)
sub_row(table, '1.4', 'Індикативні методи вимірювання', GRAY, BG_WHITE)
sub_row(table, '1.5', 'Проблеми існуючої системи моніторингу атмосферного повітря', GRAY, BG_WHITE)

# ---- Р2 ----
chapter_row(table, 'Р 2', 'Аналіз співвідношення PM₁₀/TSP', 'ok', BG_GREEN, GREEN, 'без змін', BG_GREEN)
sub_row(table, '2.1', 'Міжнародні дослідження PM₁₀/TSP', GRAY, BG_WHITE)
sub_row(table, '2.2', 'Вітчизняні дослідження', GRAY, BG_WHITE)
sub_row(table, '2.3', 'Фактори впливу на співвідношення', GRAY, BG_WHITE)
sub_row(table, '2.4', 'Обґрунтування необхідності локальних досліджень', GRAY, BG_WHITE)
sub_row(table, '2.5', 'Вплив точності визначення PM₁₀ на оцінку екологічної безпеки', GRAY, BG_WHITE)

# ---- Р3 ----
chapter_row(table, 'Р 3', 'Обґрунтування об\'єкта та схеми дослідження', 'fix', BG_ORANGE, ORANGE, '2 доповнення', BG_ORANGE)
sub_row(table, '3.1', 'Вибір міста Кривий Ріг', GRAY, BG_WHITE)
sub_row(table, '3.2', 'Аналіз мережі постів', ORANGE, BG_ORANGE, '+ три пости розміщені навмисно для діагностики зональності')
sub_row(table, '3.3', 'Обґрунтування паралельних вимірювань', ORANGE, BG_ORANGE, '+ колокаційний пост — умова однорідності середовища')
sub_row(table, '3.4', 'Вибір поста №3', GRAY, BG_WHITE)

# ---- Р4 ----
chapter_row(table, 'Р 4', 'Формування та валідація даних', 'ok', BG_GREEN, GREEN, 'без змін', BG_GREEN)
sub_row(table, '4.1', 'Джерела даних (TSP, PM₁₀)', GRAY, BG_WHITE)
sub_row(table, '4.2', 'Цифровізація даних', GRAY, BG_WHITE)
sub_row(table, '4.3', 'Синхронізація даних (алгоритм ±30 хв)', GRAY, BG_WHITE)
sub_row(table, '4.4', 'Контроль якості (QC)', GRAY, BG_WHITE)
sub_row(table, '4.5', 'Формування вибірки', GRAY, BG_WHITE)
sub_row(table, '4.6', 'Роль якості даних у забезпеченні екологічної безпеки', GRAY, BG_WHITE)

# ---- Р5 ----
chapter_row(table, 'Р 5', 'Валідація індикативних вимірювань PM₁₀', 'fix', BG_ORANGE, ORANGE, 'переписати 5.3', BG_ORANGE)
sub_row(table, '5.1', 'Обґрунтування необхідності валідації індикативних вимірювань', GRAY, BG_WHITE)
sub_row(table, '5.2', 'Комплексна валідація — трипостова (pan2, pan5, pan6)', GRAY, BG_WHITE)
sub_row(table, '5.2.1–5.2.8', 'Детальний аналіз по місяцях, метео, межі застосування', GRAY, BG_WHITE)
sub_row(table, '5.3', 'Узагальнення результатів валідації', ORANGE, BG_ORANGE, 'переписати: прилади коректні / pan5 нижчий / метод масштабується')
sub_row(table, '5.4', 'Вплив достовірності вимірювань на екологічну безпеку', GRAY, BG_WHITE)

# ---- Р6 (НОВИЙ) ----
chapter_row(table, 'Р 6', 'Валідація референтних вимірювань TSP та виявлення зон міста', 'new', BG_BLUE, BLUE, 'НОВИЙ', BG_BLUE)
sub_row(table, '6.1', 'Мережа референтних постів Гідромету Кривого Рогу', BLUE, BG_BLUE)
sub_row(table, '6.2', 'Методика порівняльного аналізу TSP між постами', BLUE, BG_BLUE)
sub_row(table, '6.3', 'Просторовий аналіз TSP — виявлення однорідності/неоднорідності', BLUE, BG_BLUE)
sub_row(table, '6.4', 'Виявлення аерозольних зон міста', BLUE, BG_BLUE)
sub_row(table, '6.5', 'Висновок: кількість зон і необхідних колокаційних точок', BLUE, BG_BLUE)

# ---- Р7 (НОВИЙ) ----
chapter_row(table, 'Р 7', 'Аналіз коефіцієнта PM₁₀/TSP по зонах', 'new', BG_BLUE, BLUE, 'НОВИЙ', BG_BLUE)
sub_row(table, '7.1', 'Методика розрахунку K', BLUE, BG_BLUE)
sub_row(table, '7.2', 'Статистичний аналіз K (медіана, σ, довірчі інтервали)', BLUE, BG_BLUE)
sub_row(table, '7.3', 'Сезонна залежність K', BLUE, BG_BLUE)
sub_row(table, '7.4', 'Метеорологічна залежність K (RH, T, вітер)', BLUE, BG_BLUE)
sub_row(table, '7.5', 'K окремо для кожної виявленої зони', BLUE, BG_BLUE)
sub_row(table, '7.6', 'Порівняння з міжнародними даними (діапазон 0.32–0.91)', BLUE, BG_BLUE)

# ---- Р8 (НОВИЙ) ----
chapter_row(table, 'Р 8', 'Розроблення RF моделей корекції та розрахунку', 'new', BG_BLUE, BLUE, 'НОВИЙ', BG_BLUE)
sub_row(table, '8.1', 'RF_M1 — модель корекції індикативних вимірювань PM₁₀', BLUE, BG_BLUE)
sub_row(table, '8.1.1', 'Архітектура: features (PM₁₀ опт + метео), target (PM₁₀ еталон з TSP)', BLUE, BG_BLUE)
sub_row(table, '8.1.2', 'Навчання на синхронізованому масиві Кривого Рогу', BLUE, BG_BLUE)
sub_row(table, '8.1.3', 'Валідація RF_M1 — невизначеність ≤ 25–30%', BLUE, BG_BLUE)
sub_row(table, '8.1.4', 'Межі застосування моделі', BLUE, BG_BLUE)
sub_row(table, '8.2', 'RF_M2 — модель локального коефіцієнта PM₁₀/TSP', BLUE, BG_BLUE)
sub_row(table, '8.2.1', 'Архітектура: features (місяць + метео + зона), target (K)', BLUE, BG_BLUE)
sub_row(table, '8.2.2', 'Навчання після RF_M1 (PM₁₀ скоригований як основа K)', BLUE, BG_BLUE)
sub_row(table, '8.2.3', 'Валідація RF_M2', BLUE, BG_BLUE)
sub_row(table, '8.2.4', 'Реалізація по зонах міста', BLUE, BG_BLUE)

# ---- Р9 (НОВИЙ) ----
chapter_row(table, 'Р 9', 'Програмний комплекс', 'new', BG_BLUE, BLUE, 'НОВИЙ', BG_BLUE)
sub_row(table, '9.1', 'Архітектура програмного комплексу (6 модулів)', BLUE, BG_BLUE)
sub_row(table, '9.2', 'Модуль 1: валідація індикативної мережі (CairCloud API)', BLUE, BG_BLUE)
sub_row(table, '9.3', 'Модуль 2: Excel форма для даних Гідромету + QC при введенні', BLUE, BG_BLUE)
sub_row(table, '9.4', 'Модуль 3: валідація TSP Гідромету + карта зон', BLUE, BG_BLUE)
sub_row(table, '9.5', 'Модуль 4: RF_M1 в роботі — автоматична корекція PM₁₀', BLUE, BG_BLUE)
sub_row(table, '9.6', 'Модуль 5: RF_M2 в роботі — PM₁₀ по мережі Гідромету', BLUE, BG_BLUE)
sub_row(table, '9.7', 'Модуль 6: AQI + перевищення + карти + звіти', BLUE, BG_BLUE)
sub_row(table, '9.8', 'Методологія масштабування на інші міста України', BLUE, BG_BLUE)

# ---- Р10 (НОВИЙ) ----
chapter_row(table, 'Р 10', 'Оцінка підвищення екологічної безпеки', 'new', BG_BLUE, BLUE, 'НОВИЙ', BG_BLUE)
sub_row(table, '10.1', 'Вплив RF_M1 на якість моніторингу (до/після корекції)', BLUE, BG_BLUE)
sub_row(table, '10.2', 'Вплив RF_M2 на охоплення мережі Гідромету', BLUE, BG_BLUE)
sub_row(table, '10.3', 'Практичні рекомендації для органів моніторингу', BLUE, BG_BLUE)
sub_row(table, '10.4', 'Відповідність Директиві (ЄС) 2024/2881', BLUE, BG_BLUE)

# Set column widths
for row in table.rows:
    row.cells[0].width = Cm(1.5)
    row.cells[1].width = Cm(12)
    row.cells[2].width = Cm(3)

doc.add_paragraph()

# Footer
p = doc.add_paragraph()
r = p.add_run('Загальні висновки   |   Список джерел   |   Додатки')
r.font.size = Pt(10)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)

separator(doc)

p = doc.add_paragraph()
r = p.add_run('E:\\TEMA — Червень 2026')
r.font.size = Pt(9)
r.font.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('E:\\TEMA\\СТРУКТУРА_ДИСЕРТАЦІЇ_v1.docx')
print('Done!')
